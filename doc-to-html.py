import webbrowser
from docx import Document
import os


def get_font_style(run):
    font_style = ''
    if run.bold:
        font_style += 'font-weight: bold; '
    if run.italic:
        font_style += 'font-style: italic; '
    if run.underline:
        font_style += 'text-decoration: underline; '
    if run.font.name:
        font_style += f'font-family: {run.font.name}; '
    if run.font.size:
        font_style += f'font-size: {run.font.size.pt}pt; '
    return font_style.strip()


def convert_docx_to_html(docx_file, html_file, output_folder):
    doc = Document(docx_file)
    image_filenames = []
    used_image_filenames = set()  # Keep track of used image filenames
    bullet_points = []

    with open(html_file, 'w', encoding='utf-8') as f:
        f.write('<html>\n<head>\n<title>Converted Document</title>\n')
        f.write('<style>\n.content {\n padding-left: 250px; \n padding-right: 250px;\n}\n</style>\n')
        f.write('</head>\n<body>\n')
        f.write('<div class="content">\n')
        
        for paragraph in doc.paragraphs:
            if paragraph.runs:
                if paragraph.style.name.startswith('List') and paragraph.text.strip():
                    bullet_points.append(paragraph.text.strip())
                else:
                    f.write('<p>')
                    for run in paragraph.runs:
                        if run.text:
                            font_style = get_font_style(run)
                            if font_style:
                                f.write('<span style="{}">{}</span>'.format(font_style, run.text))
                            else:
                                f.write(run.text)
                        else:
                            for para in doc.paragraphs:                     # hyperlinks
                                for link in para._element.xpath(".//w:hyperlink"):
                                    inner_run = link.xpath("w:r")[0]
                                    hyperlink_text = inner_run.text
                                    rId = link.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                                    hyperlink_url = doc._part.rels[rId]._target
                                    f.write(f'<a href="{hyperlink_url}">{hyperlink_text}</a>')
                            for elem in run._element.iterdescendants():     # drawing element (image)
                                if "drawing" in elem.tag:
                                    for elem2 in elem.iterdescendants():
                                        if "graphicData" in elem2.tag:
                                            for elem3 in elem2.iterdescendants():
                                                if "pic" in elem3.tag:
                                                    for elem4 in elem3.iterdescendants():
                                                        if "blip" in elem4.tag:
                                                            rId = elem4.attrib.get(
                                                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                                                            )
                                                            if rId is not None:
                                                                image_filename = f"image{rId}.png"
                                                                image_path = os.path.join(output_folder, image_filename)
                                                                if image_filename not in used_image_filenames:
                                                                    image_data = doc.part.related_parts[rId].blob
                                                                    with open(image_path, "wb") as img_file:
                                                                        img_file.write(image_data)
                                                                    image_filenames.append(image_filename)
                                                                    used_image_filenames.add(image_filename)
                                                                    f.write(f'<img src="{image_filename}" alt="Image" style="width: 100%; height: auto;" /><br />\n')
                                                                    break
                                                    break
                                            break
                                    break
                    f.write('</p>\n')
        if bullet_points:
            f.write('<ul>\n')
            for bullet_point in bullet_points:
                f.write(f'<li>{bullet_point}</li>\n')
            f.write('</ul>\n')

        f.write('</div>\n')
        f.write('</body>\n</html>')


docx_file = r"C:\Users\Shashank_Maurya\OneDrive - Dell Technologies\Documents\doc-to-html\iOSDocumentForIntuneRelease.docx"
html_file = 'output2.html'
output_folder = r"C:\Users\Shashank_Maurya\OneDrive - Dell Technologies\Documents\doc-to-html"

convert_docx_to_html(docx_file, html_file, output_folder)

webbrowser.open(html_file)
